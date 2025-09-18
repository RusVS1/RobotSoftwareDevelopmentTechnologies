import sys
import os
import openpyxl
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from openpyxl.formatting.rule import CellIsRule
from docx import Document
import subprocess

data = []

for line in sys.stdin:
    line = line.split()
    line = [line[0], int(line[1]), float(line[2])]
    data.append(line)

wb = openpyxl.Workbook()
sheet = wb.active
sheet.title = "Товары"
sheet['A1'] = "Название"
sheet['B1'] = "Количество"
sheet['C1'] = "Цена"

str_c = 2
for line in data:
    sheet['A' + str(str_c)] = line[0]
    sheet['B' + str(str_c)] = line[1]
    sheet['C' + str(str_c)] = line[2]
    str_c += 1

fill = PatternFill(start_color="FFDB8B", end_color="FFDB8B", fill_type="solid")

rule = CellIsRule(
    operator="greaterThan",
    formula=['100'],
    stopIfTrue=True,
    fill=fill
)

sheet.conditional_formatting.add('B2:B' + str(len(data)), rule)
    
wb.save("test.xlsx")

print("excel save to", os.path.abspath("test.xlsx"))

doc = Document()

doc.add_heading("Отчет сгенерирован роботом...", level=1)
doc.add_paragraph("Отчет сгенерирован роботом...")

table = doc.add_table(rows=len(data) + 1, cols=3)
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Название"
hdr_cells[1].text = "Количество"
hdr_cells[2].text = "Цена"

for i, row_data in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = str(row_data[0])
    row[1].text = str(row_data[1])
    row[2].text = str(row_data[2])

doc.save("test.docx")

print("word save to", os.path.abspath("test.docx"))

subprocess.run([
    "libreoffice",
    "--headless",
    "--convert-to",
    "pdf",
    "--outdir",
    os.path.dirname(os.path.abspath("test.docx")),
    "test.docx"
], stdout = subprocess.DEVNULL, stderr = subprocess.DEVNULL)

print("pdf save to", os.path.abspath("test.pdf"))
